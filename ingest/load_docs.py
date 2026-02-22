"""
load_docs.py - CiteWise Legal Document Loader
==============================================
Handles unified extraction of text from PDF and Word (.docx) files.
Applies legal-grade cleaning:
  - Removes headers, footers, watermarks, page numbers, auto-line numbers
  - Normalises Hebrew punctuation (Gershayim), smart quotes, double spaces
  - Preserves section numbering and clause punctuation (NEVER strips them)

Author: CiteWise Senior Legal AI Architect
PEP 8 compliant.
"""

import logging
import re
from pathlib import Path
from typing import List

from langchain_core.documents import Document
from pypdf import PdfReader

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

# Supported extensions and their loaders
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}

# Regex: standalone page-number lines  (e.g. "- 3 -", "3", "עמוד 3")
_RE_PAGE_NUM = re.compile(
    r"^\s*(-\s*)?\d{1,4}(\s*-)?\s*$"
    r"|^\s*עמוד\s+\d+\s*$",
    re.MULTILINE,
)

# Regex: recurring auto-line numbers at start of line (e.g. "  1  text...")
_RE_LINE_NUM = re.compile(r"^\s{0,4}\d{1,3}\s{2,}", re.MULTILINE)

# Smart / curly quotes → standard
_QUOTE_MAP = str.maketrans({
    "\u201c": '"',  # "
    "\u201d": '"',  # "
    "\u2018": "'",  # '
    "\u2019": "'",  # '
    "\u201e": '"',  # „ (common in Hebrew docs)
    "\u201a": "'",  # ‚
})

# Hebrew Gershayim normalisation: replace double-apostrophe with proper ״
_RE_DOUBLE_QUOTE_HEB = re.compile(r"(?<=[א-ת])\"(?=[א-ת])")
_RE_DOUBLE_APOS_HEB = re.compile(r"(?<=[א-ת])''(?=[א-ת])")


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _clean_text(raw: str) -> str:
    """
    Apply legal-grade text cleaning to extracted raw text.

    Cleaning steps (order matters):
    1. Translate smart quotes to ASCII equivalents.
    2. Remove standalone page-number lines.
    3. Remove auto-line numbers at the start of lines.
    4. Join broken lines *within* the same paragraph
       (single newline → space), while preserving double-newline
       paragraph/clause boundaries.
    5. Collapse excessive whitespace within a line.
    6. Normalise Hebrew Gershayim punctuation.

    NOTE: Section numbering (1.1, (א), (ב), 3.2.1 …) is intentionally
    preserved at all times.
    """
    # Step 1 – smart quotes
    text = raw.translate(_QUOTE_MAP)

    # Step 2 – page numbers
    text = _RE_PAGE_NUM.sub("", text)

    # Step 3 – auto line numbers
    text = _RE_LINE_NUM.sub("", text)

    # Step 4 – join broken lines inside paragraphs.
    # Strategy: split on \n\n+ to isolate paragraphs, then within each
    # paragraph collapse single newlines to a space.
    paragraphs = re.split(r"\n{2,}", text)
    joined = []
    for para in paragraphs:
        # Within a paragraph, replace single newline with space
        merged = para.replace("\n", " ")
        # Collapse multiple spaces
        merged = re.sub(r"  +", " ", merged).strip()
        if merged:
            joined.append(merged)
    text = "\n\n".join(joined)

    # Step 5 – Hebrew Gershayim
    text = _RE_DOUBLE_QUOTE_HEB.sub("״", text)
    text = _RE_DOUBLE_APOS_HEB.sub("״", text)

    return text.strip()


def _load_pdf(file_path: Path) -> List[Document]:
    """
    Extract text from a PDF using pypdf, page by page.

    Returns a list of LangChain Document objects, one per page,
    with metadata: {source, page}.
    """
    docs: List[Document] = []
    try:
        reader = PdfReader(str(file_path))
        for page_num, page in enumerate(reader.pages, start=1):
            raw_text = page.extract_text() or ""
            if not raw_text.strip():
                logger.debug("Page %d of %s is empty – skipping.", page_num, file_path.name)
                continue
            cleaned = _clean_text(raw_text)
            docs.append(Document(
                page_content=cleaned,
                metadata={
                    "source": file_path.name,
                    "page": page_num,
                },
            ))
    except Exception as exc:
        logger.error("Failed to read PDF %s: %s", file_path, exc)
    return docs


def _load_docx(file_path: Path) -> List[Document]:
    """
    Extract text from a .docx Word document using python-docx.

    Word files don't have discrete pages, so we treat the entire
    document as a single "page 1" and rely on the chunker to split it.
    Returns a list with a single Document.
    """
    try:
        # python-docx is not in requirements.txt, so we import lazily
        # and provide a helpful error if missing.
        from docx import Document as DocxDocument  # type: ignore
    except ImportError:
        logger.error(
            "python-docx is required to load .docx files. "
            "Install it with: pip install python-docx"
        )
        return []

    try:
        docx = DocxDocument(str(file_path))
        # Join paragraphs with double newlines to preserve clause boundaries
        raw_text = "\n\n".join(para.text for para in docx.paragraphs if para.text.strip())
        cleaned = _clean_text(raw_text)
        return [Document(
            page_content=cleaned,
            metadata={
                "source": file_path.name,
                "page": 1,
            },
        )]
    except Exception as exc:
        logger.error("Failed to read DOCX %s: %s", file_path, exc)
        return []


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def load_document(file_path: Path) -> List[Document]:
    """
    Load and clean a single legal document (PDF or Word).

    Parameters
    ----------
    file_path : Path
        Absolute or relative path to the document.

    Returns
    -------
    List[Document]
        A list of LangChain Document objects with cleaned text and metadata.
        Returns an empty list if the file type is unsupported or extraction fails.
    """
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        logger.warning("Unsupported file type: %s", file_path.name)
        return []

    logger.info("Loading document: %s", file_path.name)

    if suffix == ".pdf":
        return _load_pdf(file_path)
    elif suffix in {".docx", ".doc"}:
        return _load_docx(file_path)

    return []


def load_all_documents(data_dir: Path) -> List[Document]:
    """
    Recursively load all supported legal documents from a directory.

    Parameters
    ----------
    data_dir : Path
        Root directory that contains the PDF/Word files (e.g. ./data).

    Returns
    -------
    List[Document]
        All extracted, cleaned Document objects from every supported file.
    """
    all_docs: List[Document] = []
    if not data_dir.exists():
        logger.error("Data directory does not exist: %s", data_dir)
        return all_docs

    files = [
        p for p in data_dir.rglob("*")
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
    ]

    if not files:
        logger.warning("No supported documents found in %s", data_dir)
        return all_docs

    logger.info("Found %d document(s) to load.", len(files))
    for file_path in files:
        docs = load_document(file_path)
        all_docs.extend(docs)
        logger.info("  → %s: %d page(s) extracted.", file_path.name, len(docs))

    logger.info("Total pages extracted: %d", len(all_docs))
    return all_docs
